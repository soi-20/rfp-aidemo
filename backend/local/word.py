import requests
import psycopg2
import re
from docx import Document
import os
from azure.identity import ClientSecretCredential
from dotenv import load_dotenv
load_dotenv()

credentials = ClientSecretCredential(
    client_id="86dfb3ef-b0fc-4bbe-90b6-f1a11c619f8e",
    client_secret="34d8Q~B72MZ5PUCQKI5ZCR2L3cM6PB4eJ1teFbw.",
    tenant_id="1b314d3f-4df5-485a-a59b-72335826d944"
)

scopes = ["https://graph.microsoft.com/.default"]


def getDataServiceFMv2(credentials, scopes, site_id, drive_id, workbook_id, worksheet_id):
    access_token = credentials.get_token(*scopes)
    endpoint = f'https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/items/{workbook_id}/workbook/worksheets/{worksheet_id}/usedRange'
    headers = {
        'Authorization': f'Bearer {access_token.token}'
    }
    response = requests.get(endpoint, headers=headers)
    if response.status_code == 200:
        worksheet_data = response.json()
        final_data = []
        data = worksheet_data['values']
        for i, row in enumerate(data):
            num = row[3] 
            pattern = re.compile(r'^\d+\.\d+\.\d+$')
            if pattern.match(str(num)):
                final_data.append({
                    "row_number": i + 1,
                    "data": row
                })
        return final_data
    else:
        print(f"Failed to retrieve worksheet data: {response.status_code}")

conn = psycopg2.connect(
    user= f"{os.environ.get('POSTGRES_USER')}", 
    password=f"{os.environ.get('POSTGRES_PASSWORD')}", 
    host=f"{os.environ.get('POSTGRES_HOST')}", 
    port=f"{os.environ.get('POSTGRES_PORT')}", 
    database=f"{os.environ.get('POSTGRES_DATABASE')}"
)
cur = conn.cursor()

site_name = "ServiceFm-v2"
cur.execute(f"SELECT * FROM site WHERE name = '{site_name}'")
site = cur.fetchone()
print(site)
data = getDataServiceFMv2(credentials, scopes, site_id = site[2], drive_id = site[3], workbook_id = site[4], worksheet_id = site[5])

document = Document()
document.add_heading('ServiceFM-v2', 0)


for row in data:
    values = row['data']
    if values[6]  != "Text":
        continue

    document.add_heading(f'{values[3]} {values[4]}', level=4)
    document.add_paragraph(f'{values[5]}')
    document.add_heading('Response', level=5)
    document.add_paragraph(values[8])


# document.save('C:/Users/anubh/Desktop/ServiceFM-v2.docx')



